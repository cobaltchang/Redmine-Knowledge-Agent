"""Attachment processors using Factory Pattern.

Each processor handles a specific type of attachment and extracts content.
"""

from __future__ import annotations

import mimetypes
from abc import ABC, abstractmethod
from typing import TYPE_CHECKING, Any, Protocol, runtime_checkable

from .models import ExtractedContent, ProcessingMethod

if TYPE_CHECKING:
    from pathlib import Path

    from .config import ProcessingConfig

# Optional dependencies - declare placeholders then attempt imports (avoid redefinition warnings)
Image: Any = None
pytesseract: Any = None
Document: Any = None
openpyxl: Any = None

try:
    from PIL import Image as _PIL_Image

    Image = _PIL_Image
except ImportError:  # pragma: no cover
    Image = None

try:
    import pytesseract as _pytesseract  # type: ignore

    pytesseract = _pytesseract
except ImportError:  # pragma: no cover
    pytesseract = None

fitz: Any = None

try:
    import fitz as _fitz  # type: ignore

    fitz = _fitz
except ImportError:  # pragma: no cover
    fitz = None

try:
    from docx import Document as _DocxDocument

    Document = _DocxDocument
except ImportError:  # pragma: no cover
    Document = None

try:
    import openpyxl as _openpyxl  # type: ignore

    openpyxl = _openpyxl
except ImportError:  # pragma: no cover
    openpyxl = None


@runtime_checkable
class AttachmentProcessor(Protocol):
    """Protocol for attachment processors."""

    def process(self, file_path: Path) -> ExtractedContent:
        """Process an attachment and extract content.

        Args:
            file_path: Path to the downloaded attachment file.

        Returns:
            ExtractedContent with extracted text and metadata.

        """
        ...  # pragma: no cover

    @property
    def supported_types(self) -> list[str]:
        """List of supported MIME types."""
        ...  # pragma: no cover


class BaseProcessor(ABC):
    """Base class for attachment processors."""

    def __init__(self, config: ProcessingConfig | None = None) -> None:
        """Initialize the processor.

        Args:
            config: Processing configuration.

        """
        self.config = config

    @abstractmethod
    def process(self, file_path: Path) -> ExtractedContent:
        """Process an attachment and extract content."""
        ...  # pragma: no cover

    @property
    @abstractmethod
    def supported_types(self) -> list[str]:
        """List of supported MIME types."""
        ...  # pragma: no cover

    def _create_error_result(self, error: str) -> ExtractedContent:
        """Create an error result.

        Args:
            error: Error message.

        Returns:
            ExtractedContent with error information.

        """
        return ExtractedContent(
            text="",
            metadata={},
            processing_method=ProcessingMethod.FALLBACK,
            error=error,
        )


class ImageProcessor(BaseProcessor):
    """Processor for image files using OCR."""

    @property
    def supported_types(self) -> list[str]:
        """Return list of supported MIME types for image processing."""
        return [
            "image/png",
            "image/jpeg",
            "image/jpg",
            "image/gif",
            "image/bmp",
            "image/tiff",
            "image/webp",
        ]

    def process(self, file_path: Path) -> ExtractedContent:
        """Process an image file using OCR.

        Args:
            file_path: Path to the image file.

        Returns:
            ExtractedContent with OCR-extracted text.

        """
        if not file_path.exists():
            return self._create_error_result(f"File not found: {file_path}")

        if pytesseract is None or Image is None:
            return self._create_error_result(
                "OCR dependencies not available: pytesseract or PIL not installed",
            )

        try:
            with Image.open(file_path) as img:
                # Use a converted copy when needed to avoid assigning incompatible types
                if getattr(img, "mode", None) in ("RGBA", "P"):
                    img_to_process = img.convert("RGB")
                else:
                    img_to_process = img

                text = pytesseract.image_to_string(img_to_process, lang="eng+chi_tra")

            return ExtractedContent(
                text=text.strip(),
                metadata={
                    "filename": file_path.name,
                    "size": file_path.stat().st_size,
                },
                processing_method=ProcessingMethod.OCR,
            )
        except (OSError, ValueError, RuntimeError) as e:
            return self._create_error_result(f"OCR failed: {e}")


class PdfProcessor(BaseProcessor):
    """Processor for PDF files."""

    @property
    def supported_types(self) -> list[str]:
        """Return list of supported MIME types for PDF processing."""
        return ["application/pdf"]

    def process(self, file_path: Path) -> ExtractedContent:
        """Process a PDF file and extract text.

        Args:
            file_path: Path to the PDF file.

        Returns:
            ExtractedContent with extracted text.

        """
        if not file_path.exists():
            return self._create_error_result(f"File not found: {file_path}")

        try:
            # Prefer module-level `fitz` (can be patched in tests), fallback to dynamic import
            fitz_mod = fitz
            if fitz_mod is None:
                return self._create_error_result(
                    "PDF dependencies not available: PyMuPDF not installed",
                )

            text_parts: list[str] = []
            page_count = 0

            with fitz_mod.open(str(file_path)) as doc:
                page_count = len(doc)
                for page in doc:  # pragma: no branch
                    page_text = page.get_text()
                    if page_text.strip():  # pragma: no branch
                        text_parts.append(page_text)

            return ExtractedContent(
                text="\n\n".join(text_parts),
                metadata={
                    "filename": file_path.name,
                    "page_count": page_count,
                    "size": file_path.stat().st_size,
                },
                processing_method=ProcessingMethod.TEXT_EXTRACT,
            )
        except (OSError, ValueError, RuntimeError) as e:
            return self._create_error_result(f"PDF extraction failed: {e}")


class DocxProcessor(BaseProcessor):
    """Processor for Word documents."""

    @property
    def supported_types(self) -> list[str]:
        """Return list of supported MIME types for DOCX processing."""
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]

    def process(self, file_path: Path) -> ExtractedContent:
        """Process a Word document and extract text.

        Args:
            file_path: Path to the Word document.

        Returns:
            ExtractedContent with extracted text.

        """
        if not file_path.exists():
            return self._create_error_result(f"File not found: {file_path}")

        if Document is None:
            return self._create_error_result(
                "DOCX dependencies not available: python-docx not installed",
            )

        try:
            doc = Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract text from tables
            table_texts: list[str] = []
            for table in doc.tables:  # pragma: no branch
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:  # pragma: no branch
                        table_texts.append(row_text)

            all_text = "\n\n".join(paragraphs)
            if table_texts:
                all_text += "\n\n### Tables\n\n" + "\n".join(table_texts)

            return ExtractedContent(
                text=all_text,
                metadata={
                    "filename": file_path.name,
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables),
                    "size": file_path.stat().st_size,
                },
                processing_method=ProcessingMethod.TEXT_EXTRACT,
            )
        except (OSError, ValueError, RuntimeError) as e:
            return self._create_error_result(f"DOCX extraction failed: {e}")


class SpreadsheetProcessor(BaseProcessor):
    """Processor for spreadsheet files."""

    @property
    def supported_types(self) -> list[str]:
        """Return list of supported MIME types for spreadsheet processing."""
        return [
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
            "text/csv",
        ]

    def process(self, file_path: Path) -> ExtractedContent:
        """Process a spreadsheet and convert to Markdown table.

        Args:
            file_path: Path to the spreadsheet file.

        Returns:
            ExtractedContent with Markdown table representation.

        """
        if not file_path.exists():
            return self._create_error_result(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()

        try:
            if suffix == ".csv":
                return self._process_csv(file_path)
            if openpyxl is None:
                return self._create_error_result(
                    "Spreadsheet dependencies not available: openpyxl not installed",
                )
            return self._process_excel(file_path)
        except (OSError, ValueError, RuntimeError) as e:
            return self._create_error_result(f"Spreadsheet extraction failed: {e}")

    def _process_csv(self, file_path: Path) -> ExtractedContent:
        """Process a CSV file."""
        import csv  # noqa: PLC0415 - deferred import for optional dependency

        rows: list[list[str]] = []
        with file_path.open(encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            rows = list(reader)

        if not rows:
            return ExtractedContent(
                text="(Empty spreadsheet)",
                metadata={"filename": file_path.name, "row_count": 0},
                processing_method=ProcessingMethod.TEXT_EXTRACT,
            )

        md_table = self._rows_to_markdown(rows)
        return ExtractedContent(
            text=md_table,
            metadata={
                "filename": file_path.name,
                "row_count": len(rows),
                "size": file_path.stat().st_size,
            },
            processing_method=ProcessingMethod.TEXT_EXTRACT,
        )

    def _process_excel(self, file_path: Path) -> ExtractedContent:
        """Process an Excel file."""
        wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
        all_sheets: list[str] = []
        total_rows = 0

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows: list[list[str]] = []
            for row in sheet.iter_rows(values_only=True):  # pragma: no branch
                str_row = [str(cell) if cell is not None else "" for cell in row]
                if any(str_row):  # Skip completely empty rows  # pragma: no branch
                    rows.append(str_row)

            if rows:  # pragma: no branch
                total_rows += len(rows)
                md_table = self._rows_to_markdown(rows)
                all_sheets.append(f"### {sheet_name}\n\n{md_table}")

        wb.close()

        return ExtractedContent(
            text="\n\n".join(all_sheets) if all_sheets else "(Empty spreadsheet)",
            metadata={
                "filename": file_path.name,
                "sheet_count": len(wb.sheetnames),
                "total_rows": total_rows,
                "size": file_path.stat().st_size,
            },
            processing_method=ProcessingMethod.TEXT_EXTRACT,
        )

    def _rows_to_markdown(self, rows: list[list[str]]) -> str:
        """Convert rows to Markdown table format."""
        if not rows:
            return ""

        # Use first row as header
        header = rows[0]
        col_count = len(header)

        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * col_count) + " |",
        ]

        for row in rows[1:]:
            # Pad row if necessary
            padded = row + [""] * (col_count - len(row))
            lines.append("| " + " | ".join(padded[:col_count]) + " |")

        return "\n".join(lines)


class FallbackProcessor(BaseProcessor):
    """Fallback processor for unsupported file types."""

    @property
    def supported_types(self) -> list[str]:
        """Return wildcard MIME type pattern (matches all files)."""
        return ["*/*"]  # Matches everything

    def process(self, file_path: Path) -> ExtractedContent:
        """Create a fallback result with file metadata.

        Args:
            file_path: Path to the file.

        Returns:
            ExtractedContent with only metadata.

        """
        if not file_path.exists():
            return self._create_error_result(f"File not found: {file_path}")

        mime_type, _ = mimetypes.guess_type(str(file_path))

        return ExtractedContent(
            text=f"(Binary file: {file_path.name}, type: {mime_type or 'unknown'})",
            metadata={
                "filename": file_path.name,
                "mime_type": mime_type,
                "size": file_path.stat().st_size,
            },
            processing_method=ProcessingMethod.FALLBACK,
        )


class ProcessorFactory:
    """Factory for creating attachment processors."""

    def __init__(self, config: ProcessingConfig | None = None) -> None:
        """Initialize the factory.

        Args:
            config: Processing configuration.

        """
        self.config = config
        self._processors: dict[str, BaseProcessor] = {}
        self._register_default_processors()

    def _register_default_processors(self) -> None:
        """Register default processors."""
        processors = [
            ImageProcessor(self.config),
            PdfProcessor(self.config),
            DocxProcessor(self.config),
            SpreadsheetProcessor(self.config),
        ]

        for processor in processors:
            for mime_type in processor.supported_types:
                self._processors[mime_type] = processor

    def register_processor(self, mime_type: str, processor: BaseProcessor) -> None:
        """Register a custom processor for a MIME type.

        Args:
            mime_type: The MIME type to handle.
            processor: The processor instance.

        """
        self._processors[mime_type] = processor

    def get_processor(self, mime_type: str, filename: str | None = None) -> BaseProcessor:
        """Get a processor for the given MIME type.

        Args:
            mime_type: The MIME type of the attachment.
            filename: Optional filename for additional type detection.

        Returns:
            An appropriate processor for the file type.

        """
        # Direct match
        if mime_type in self._processors:
            return self._processors[mime_type]

        # Try to guess from filename
        if filename:
            guessed_type, _ = mimetypes.guess_type(filename)
            if guessed_type and guessed_type in self._processors:
                return self._processors[guessed_type]

        # Fallback
        return FallbackProcessor(self.config)

    def process_file(self, file_path: Path, mime_type: str | None = None) -> ExtractedContent:
        """Process a file using the appropriate processor.

        Args:
            file_path: Path to the file.
            mime_type: Optional MIME type (will be guessed if not provided).

        Returns:
            ExtractedContent with processed content.

        """
        if mime_type is None:
            mime_type, _ = mimetypes.guess_type(str(file_path))
            mime_type = mime_type or "application/octet-stream"

        processor = self.get_processor(mime_type, file_path.name)
        return processor.process(file_path)
